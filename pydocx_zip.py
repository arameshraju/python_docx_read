import os
import docx
import zipfile
import xml.etree.ElementTree as ET

def get_docx_metadata(file_path):
    doc = docx.Document(file_path)
    properties = doc.core_properties
    word_count = sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)
    
    # Extract page count from docProps/app.xml
    page_count = None
    with zipfile.ZipFile(file_path, 'r') as docx_zip:
        if 'docProps/app.xml' in docx_zip.namelist():
            with docx_zip.open('docProps/app.xml') as app_xml:
                tree = ET.parse(app_xml)
                root = tree.getroot()
                page_count = root.find('{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}Pages').text
    
    # Extract titles and headings
    titles_and_headings = []
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            titles_and_headings.append(paragraph.text)
    
    metadata = {
        "Title": properties.title,
        "Author": properties.author,
        "Word Count": word_count,
        "Page Count": page_count,
        "Titles and Headings": titles_and_headings
    }
    
    return metadata

def main():
    docx_folder = 'docx'
    with open('pydocx_zip.md', 'w') as md_file:
        for filename in os.listdir(docx_folder):
            if filename.endswith('.docx'):
                file_path = os.path.join(docx_folder, filename)
                metadata = get_docx_metadata(file_path)
                md_file.write(f"## Metadata for {filename}\n")
                for key, value in metadata.items():
                    md_file.write(f"**{key}**: {value}\n")
                md_file.write("\n")

if __name__ == "__main__":
    main()