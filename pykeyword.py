import os
import docx
import zipfile
from keybert import KeyBERT
import xml.etree.ElementTree as ET

def get_docx_metadata(file_path):
    doc = docx.Document(file_path)
    properties = doc.core_properties
    word_count = sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)
    
    # Extract page count from docProps/app.xml
    page_count = None
    with zipfile.ZipFile(file_path, 'r') as docx_zip:
        if 'docProps/app.xml' in docx_zip.namelist():
            with docx_zip.open('docProps/app.xml') as app_xml:
                tree = ET.parse(app_xml)
                root = tree.getroot()
                page_count = root.find('{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}Pages').text
    
    # Extract titles and headings using keybert
    titles_and_headings = []
    # read the document text
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + ' '
    # extract keywords
    model = KeyBERT('distilbert-base-nli-mean-tokens')
    keywords = model.extract_keywords(text, keyphrase_ngram_range=(1, 1), stop_words='english', use_maxsum=True, nr_candidates=20, top_n=5)
    
    metadata = {
        "Title": properties.title,
        "Author": properties.author,
        "Word Count": word_count,
        "Page Count": page_count,
        "Titles and Headings": keywords
    }
    
    return metadata

def main():
    docx_folder = 'docx'
    with open('keywords.md', 'w') as md_file:
        for filename in os.listdir(docx_folder):
            if filename.endswith('.docx'):
                file_path = os.path.join(docx_folder, filename)
                metadata = get_docx_metadata(file_path)
                md_file.write(f"## Metadata for {filename}\n")
                for key, value in metadata.items():
                    md_file.write(f"**{key}**: {value}\n")
                md_file.write("\n")

if __name__ == "__main__":
    main()