import os
import docx

def get_docx_metadata(file_path):
    doc = docx.Document(file_path)
    properties = doc.core_properties
    word_count = sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)
    # Note: docx does not provide a direct way to get the number of pages
    # This is a placeholder for the number of pages
    page_count = "N/A"
   
    # Extract titles and headings
    titles_and_headings = []
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            titles_and_headings.append(paragraph.text)
    
    metadata = {
        "Title": properties.title,
        "Author": properties.author,
        "Word Count": word_count,
        "Page Count": page_count,
        "Titles and Headings": titles_and_headings
    }
    
    return metadata

def main():
    docx_folder = 'docx'
    with open('pydocx.md', 'w') as md_file:
        for filename in os.listdir(docx_folder):
            if filename.endswith('.docx'):
                file_path = os.path.join(docx_folder, filename)
                metadata = get_docx_metadata(file_path)
                md_file.write(f"## Metadata for {filename}\n")
                for key, value in metadata.items():
                    md_file.write(f"**{key}**: {value}\n")
                md_file.write("\n")

if __name__ == "__main__":
    main()